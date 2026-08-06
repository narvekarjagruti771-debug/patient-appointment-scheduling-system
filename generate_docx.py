# /// script
# dependencies = [
#   "python-docx",
#   "pymupdf"
# ]
# ///

import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz

def extract_pdf_pages(pdf_path, max_pages=4):
    print(f"Extracting first {max_pages} pages from {pdf_path} as images...")
    doc_pdf = fitz.open(pdf_path)
    img_paths = []
    for page_num in range(min(max_pages, len(doc_pdf))):
        page = doc_pdf.load_page(page_num)
        pix = page.get_pixmap(dpi=200) # High-quality rendering
        img_path = f"temp_page_{page_num}.png"
        pix.save(img_path)
        img_paths.append(img_path)
    return img_paths

def add_p(doc, text, bold=False, italic=False, style='Normal', space_after=4, space_before=0, keep_together=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    if keep_together:
        p.paragraph_format.keep_together = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet_point_docx(doc, bold_part, text_part):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    run_bold = p.add_run(f"{bold_part}: ")
    run_bold.font.name = 'Arial'
    run_bold.font.size = Pt(11)
    run_bold.bold = True
    
    run_text = p.add_run(text_part)
    run_text.font.name = 'Arial'
    run_text.font.size = Pt(11)
    return p

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42) # Slate-900
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(3, 105, 161) # Sky-700
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11.5)
    run.bold = True
    run.font.color.rgb = RGBColor(51, 65, 85) # Slate-700
    return p

def add_diagram_docx(doc, image_path, caption_text, width_inches=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(8)
    run_cap = p_cap.add_run(caption_text)
    run_cap.font.name = 'Arial'
    run_cap.font.size = Pt(9.5)
    run_cap.italic = True
    run_cap.font.color.rgb = RGBColor(71, 85, 105)
    return p

def add_dictionary_table(doc, title, rows):
    add_p(doc, f"Table Structure: {title}", bold=True, space_after=2, space_before=4)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Format Header Row
    hdr_cells = table.rows[0].cells
    hdr_titles = ["Field", "Type", "Key", "Constraints / Purpose"]
    for cell, t in zip(hdr_cells, hdr_titles):
        cell.text = t
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(9.5)
        p.runs[0].font.name = 'Arial'
        p.runs[0].font.color.rgb = RGBColor(15, 23, 42)
        
    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = text
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.name = 'Arial'
            if idx == 2: # Key column center align
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() # space after table

def add_code_block_docx(doc, file_name, code):
    add_p(doc, f"[File: {file_name}]", bold=True, space_after=2, space_before=4)
    
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    
    # Gray background fill
    shading_elm = docx.oxml.parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(docx.oxml.ns.nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    
    doc.add_paragraph() # space after table

def build_docx_report():
    print("Initializing Word Document...")
    doc = docx.Document()
    
    # 1. Preliminary 4 Pages (Cover page, etc.)
    pdf_path = r"C:\Users\HP\Downloads\reppppppppp.pdf"
    if os.path.exists(pdf_path):
        img_paths = extract_pdf_pages(pdf_path, max_pages=4)
        
        # We configure a section with 0 margins for full-bleed images
        section = doc.sections[0]
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        
        for idx, img_path in enumerate(img_paths):
            doc.add_picture(img_path, width=Inches(8.27), height=Inches(11.69))
            # Delete temp image after adding
            try:
                os.remove(img_path)
            except:
                pass
            if idx < len(img_paths) - 1:
                doc.add_page_break()
                
        # Start a new section for the text report with uniform 20mm (0.787 in) margins
        report_section = doc.add_section()
        report_section.top_margin = Inches(0.787)
        report_section.bottom_margin = Inches(0.787)
        report_section.left_margin = Inches(0.787)
        report_section.right_margin = Inches(0.787)
        report_section.page_width = Inches(8.27)
        report_section.page_height = Inches(11.69)
    else:
        print(f"Warning: {pdf_path} not found. Starting report directly.")
        # Setup margins for default section instead
        section = doc.sections[0]
        section.top_margin = Inches(0.787)
        section.bottom_margin = Inches(0.787)
        section.left_margin = Inches(0.787)
        section.right_margin = Inches(0.787)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        
    # Set heading styles custom colors
    # Chapter 1: Introduction
    add_heading_1(doc, "Chapter 1: Introduction")
    doc.add_page_break() # Start chapter body on fresh page or keep together
    
    add_heading_2(doc, "1.1 Introduction")
    add_p(doc, 
        "Modern medical facilities and clinics face significant operational hurdles when managing patient appointments, consultation slots, and physician availability using manual, paper-based records. Manual processes frequently lead to scheduling overlapping appointments (double-bookings), patient queue blockages, high administrator stress, and structural record-keeping errors. This lack of coordination deteriorates the patient experience and reduces healthcare delivery efficiency."
    )
    add_p(doc,
        "To address these operational deficits, this Database Management System (DBMS) project introduces 'MediFlow', a dedicated Patient Appointment Scheduling System. MediFlow utilizes a three-tier software model that encapsulates clinic logical schemas directly inside a secure Relational Database Management System. By mapping real-world medical entities (patients, doctors, logs, and summaries) into a normalized MySQL schema, clinic admins can automate appointment queues, track historical records, and ensure absolute consistency."
    )
    
    add_heading_2(doc, "1.2 Objective")
    add_p(doc,
        "The fundamental objective of this database application is to deliver a reliable, automated, and conflict-free tool for medical schedule management. Key database-specific design targets include:"
    )
    add_bullet_point_docx(doc, "Integrity and Non-Overlapping Scheduling", "Enforce relational foreign keys and transaction boundaries so that patients and doctors cannot double-book the same slot.")
    add_bullet_point_docx(doc, "Encapsulation of Query Joins", "Create an SQL Virtual View (upcoming_appointments_view) to handle multi-table joins inside the database engine, returning a clean upcoming schedule to APIs.")
    add_bullet_point_docx(doc, "Automated Security Audit Trails", "Use database triggers to intercept appointment changes (inserts and status updates) and write immutable log audits directly to an independent logs table.")
    add_bullet_point_docx(doc, "Row-by-Row Cursor Processing", "Utilize a procedural SQL database cursor to loop through daily appointments on demand, compile a textual summary, and save it in a reporting structure.")
    add_bullet_point_docx(doc, "Modern Architecture", "Build an easy-to-use modern web interface communicating with a fast asynchronous Node.js Express API server and a MySQL database backend.")
    
    add_heading_2(doc, "1.3 Module Breakdown")
    add_p(doc, "The architecture of MediFlow is split into five distinct functional modules:")
    add_bullet_point_docx(doc, "Patient Registration Module", "Enables the registration of new patients. Validates emails to prevent duplicates and writes patient profiles (name, DOB, gender, phone, email) to the database.")
    add_bullet_point_docx(doc, "Doctor Specialty Lookup Module", "Maintains records of available specialist doctors. Allows frontend lookups to filter slots and query matching specialist consultants.")
    add_bullet_point_docx(doc, "Appointment Booking Engine", "Manages scheduling slots. Feeds new appointment rows into MySQL while checking database constraint requirements.")
    add_bullet_point_docx(doc, "Automated Audit Module", "Contains database-level triggers (after_appointment_insert, after_appointment_update) running completely independent of application code to log scheduling actions.")
    add_bullet_point_docx(doc, "Daily Report Compilation Module", "Implements the stored procedure generate_daily_report(date). Activates a cursor query that crawls specific dates, merges patient/doctor descriptions, and outputs administrative reports.")

    # Chapter 2: Survey of Technologies
    doc.add_page_break()
    add_heading_1(doc, "Chapter 2: Survey of Technologies")
    
    add_heading_2(doc, "2.1 Software Description")
    add_p(doc,
        "The software architecture follows a modular Three-Tier client-server pattern. The client layer runs directly inside any standard web browser, sending AJAX fetch commands. The middle tier runs a lightweight Node.js Express server to handle routing, input validation, and connection pool query dispatching. The storage tier uses MySQL to execute transactional data procedures. By cleanly decoupling client rendering, API controllers, and database storage, the application remains scalable and highly maintainable."
    )
    
    add_heading_2(doc, "2.2 Languages & Architectures")
    
    add_heading_3(doc, "2.2.1 HTML (HyperText Markup Language)")
    add_p(doc,
        "HTML5 is utilized to layout the application layout. The layout features a Sidebar Navigation panel, an active Status indicator checking database connectivity, a Patient Portal (registration modal, selector list, booking form), a Doctor Queue panel (for status management), and a Developer Lab interface exposing raw table counts, live database trigger log captures, and stored procedure execution outputs. All forms include HTML5 constraints (required fields, date range limits) to prevent malformed data payloads."
    )
    
    add_heading_3(doc, "2.2.2 CSS (Cascading Style Sheets)")
    add_p(doc,
        "The visual aesthetics use premium CSS3 techniques, rendering a state-of-the-art administrative dashboard. The visual theme uses a sleek Dark Mode scheme with deep purple, indigo, and teal neon gradients. Glassmorphism is implemented via CSS backdrop-filters, rendering translucent cards. Interaction feedback is provided through CSS micro-animations on hover states, pulsing indicators, and sliding tabs. Layouts are completely responsive, adjusting cleanly using CSS Flexbox and CSS Grid frameworks."
    )
    
    add_heading_3(doc, "2.2.3 PHP (Traditional Stack Survey)")
    add_p(doc,
        "Traditionally, academic DBMS projects are structured around PHP, a server-side scripting language running on Apache in XAMPP. In a standard PHP model, the server processes SQL queries sequentially via standard drivers (PDO or mysqli) and injects variables directly into HTML templates before responding. While PHP remains a robust and reliable stack, its synchronous model blocks threads for long-running database requests. In our modern implementation, Node.js replaces PHP as the application layer, running an asynchronous, event-driven event loop that uses non-blocking database pools to handle thousands of concurrent queries without blocking."
    )
    
    add_heading_3(doc, "2.2.4 MySQL Database")
    add_p(doc,
        "MySQL is an open-source Relational Database Management System (RDBMS) configured using the transactional InnoDB storage engine. MySQL executes physical database operations. It maintains strict constraints: it prevents orphan records using foreign keys with cascading deletions, exposes a queryable relational view to combine tables efficiently, runs row-level triggers to enforce data audits, and executes stored procedures that loop through active sets using database cursors. The database connects to Node.js through a persistent pool of reusable connections, eliminating connection-startup latency."
    )

    # Chapter 3: Requirements and Analysis
    doc.add_page_break()
    add_heading_1(doc, "Chapter 3: Requirements and Analysis")
    
    add_heading_2(doc, "3.1 Requirement Specification")
    add_p(doc,
        "To ensure the system satisfies both operational clinic workflows and database design principles, we specify requirements as follows:"
    )
    add_bullet_point_docx(doc, "Functional Requirements", "Patients must register with valid emails; booking engines must map patient/doctor foreign keys to target slots; system must record logs dynamically; admins must run daily cursor reports on any date.")
    add_bullet_point_docx(doc, "Non-Functional Requirements - Integrity", "The database must prevent database inconsistencies using foreign key constraints and cascading deletes.")
    add_bullet_point_docx(doc, "Non-Functional Requirements - Performance", "Index structures and virtual SQL views must optimize lookups, returning results under 50ms.")
    add_bullet_point_docx(doc, "Non-Functional Requirements - Security", "Triggers must enforce audit logs automatically, preventing admins or developers from editing appointment details without an immutable trail.")
    
    add_heading_2(doc, "3.2 Hardware and Software Requirements")
    add_bullet_point_docx(doc, "Development OS", "Windows 10 / Windows 11 (x64 architecture)")
    add_bullet_point_docx(doc, "Processor & RAM", "Intel Core i3/i5 or AMD Ryzen 3/5, 4 GB RAM minimum (8 GB recommended)")
    add_bullet_point_docx(doc, "Database Server", "MySQL Community Server 8.0.x (packaged with XAMPP 8.1.x+) running on port 3307")
    add_bullet_point_docx(doc, "Application Runtime", "Node.js v18.x or newer with npm package manager")
    add_bullet_point_docx(doc, "Web Browser", "Modern evergreen browser with JavaScript enabled (e.g., Google Chrome)")
    
    doc.add_page_break()
    add_heading_2(doc, "3.3 Data Flow Diagrams (DFD)")
    add_p(doc,
        "Data Flow Diagrams model the movement of data elements through the clinic scheduling system, mapping processes, data stores, external entities, and data flows."
    )
    
    add_heading_3(doc, "3.3.1 Level 0 DFD (Context Diagram)")
    add_p(doc,
        "The Context Diagram establishes system boundaries. It shows the main system (MediFlow DBMS) interacting with two external entities: Patients (who register and book slots) and Doctors/Admins (who configure schedules and run reports)."
    )
    add_diagram_docx(doc, "dfd_level_0.png", "Figure 3.1: Level 0 DFD - Context Diagram", width_inches=6.0)
    
    doc.add_page_break()
    add_heading_3(doc, "3.3.2 Level 1 DFD (Functional Processes)")
    add_p(doc,
        "The Level 1 DFD decomposes the system into functional process nodes. It shows how patient details flow into process 1.0 (Patient Registration) to write to the 'patients' data store, how slots map to process 3.0 (Book Appointment) connecting to 'doctors' and 'appointments' stores, how updates run through process 4.0 to generate audit logs, and how process 5.0 compiles daily agenda reports."
    )
    add_diagram_docx(doc, "dfd_level_1.png", "Figure 3.2: Level 1 DFD - Functional Processes and Data Stores", width_inches=6.0)
    
    doc.add_page_break()
    add_heading_3(doc, "3.3.3 Level 2 DFD (DBMS Internal Processing)")
    add_p(doc,
        "The Level 2 DFD details the internal database execution sequences. It visualizes the flow when SQL statements hit the DBMS: how inserts/updates trigger 'after_appointment_insert' and 'after_appointment_update' to write automatically to 'appointment_logs', and how executing the Stored Procedure initializes a cursor loop on the joined tables to output summary blocks to 'daily_reports'."
    )
    add_diagram_docx(doc, "dfd_level_2.png", "Figure 3.3: Level 2 DFD - Internal Database Event Processing", width_inches=6.0)
    
    doc.add_page_break()
    add_heading_2(doc, "3.4 Data Dictionary")
    add_p(doc,
        "The Data Dictionary defines metadata descriptions, types, keys, and constraint rules for the database schemas:"
    )
    
    add_dictionary_table(doc, "patients", [
        ("patient_id", "INT", "PK", "Primary Key, Auto increment, uniquely identifies each patient."),
        ("name", "VARCHAR(100)", "-", "Not Null, patient's full name."),
        ("email", "VARCHAR(100)", "U", "Unique, Not Null, patient's email address."),
        ("phone", "VARCHAR(15)", "-", "Not Null, phone number for contact."),
        ("dob", "DATE", "-", "Not Null, date of birth."),
        ("gender", "ENUM('Male','Female','Other')", "-", "Not Null, gender classification."),
        ("created_at", "TIMESTAMP", "-", "Default current timestamp, record creation date.")
    ])
    
    add_dictionary_table(doc, "doctors", [
        ("doctor_id", "INT", "PK", "Primary Key, Auto increment, uniquely identifies each doctor."),
        ("name", "VARCHAR(100)", "-", "Not Null, doctor's full name."),
        ("specialization", "VARCHAR(100)", "-", "Not Null, medical specialty field."),
        ("email", "VARCHAR(100)", "U", "Unique, Not Null, professional contact email."),
        ("phone", "VARCHAR(15)", "-", "Not Null, contact number."),
        ("created_at", "TIMESTAMP", "-", "Default current timestamp, record creation date.")
    ])
    
    add_dictionary_table(doc, "appointments", [
        ("appointment_id", "INT", "PK", "Primary Key, Auto increment, appointment transaction ID."),
        ("patient_id", "INT", "FK", "Foreign Key references patients(patient_id) with cascading delete."),
        ("doctor_id", "INT", "FK", "Foreign Key references doctors(doctor_id) with cascading delete."),
        ("appointment_date", "DATE", "-", "Not Null, target date of consultation."),
        ("appointment_time", "TIME", "-", "Not Null, target slot time."),
        ("status", "ENUM('Scheduled','Completed','Cancelled')", "-", "Default 'Scheduled', status indicator."),
        ("reason", "VARCHAR(255)", "-", "Not Null, patient's complaints or notes."),
        ("created_at", "TIMESTAMP", "-", "Default current timestamp, booking timestamp.")
    ])
    
    add_dictionary_table(doc, "appointment_logs", [
        ("log_id", "INT", "PK", "Primary Key, Auto increment, logs audit identifier."),
        ("appointment_id", "INT", "-", "Not Null, links log to target appointment ID."),
        ("action_type", "ENUM('CREATE','UPDATE','DELETE')", "-", "Not Null, SQL operation type caught by trigger."),
        ("old_status", "VARCHAR(20)", "-", "Nullable, previous status of updated booking."),
        ("new_status", "VARCHAR(20)", "-", "Not Null, status after trigger execution."),
        ("log_timestamp", "TIMESTAMP", "-", "Default current timestamp, log execution date/time."),
        ("description", "TEXT", "-", "Constructed text detailing the exact data changes.")
    ])
    
    add_dictionary_table(doc, "daily_reports", [
        ("report_id", "INT", "PK", "Primary Key, Auto increment, daily summary report ID."),
        ("report_date", "DATE", "U", "Unique, Not Null, target report agenda date."),
        ("report_summary", "TEXT", "-", "Not Null, aggregated text generated by cursor procedure."),
        ("generated_at", "TIMESTAMP", "-", "Auto-timestamps compilation date/time updates.")
    ])
    
    doc.add_page_break()
    add_heading_2(doc, "3.5 Entity-Relationship (ER) Diagram")
    add_p(doc,
        "The Entity-Relationship Diagram maps the structural relationships between logical entities in the MediFlow system. The schema enforces strict referential mappings:"
    )
    add_bullet_point_docx(doc, "Patient to Appointment", "One patient can book multiple appointments (1-to-Many), mapped via patient_id foreign key constraint.")
    add_bullet_point_docx(doc, "Doctor to Appointment", "One doctor can consult for multiple appointments (1-to-Many), mapped via doctor_id foreign key constraint.")
    add_bullet_point_docx(doc, "Appointment to Log Record", "One appointment modification event generates logs (1-to-Many history tracking) written automatically via MySQL triggers.")
    add_bullet_point_docx(doc, "Daily Report", "An independent report compilation object mapping report summaries to unique dates.")
    add_diagram_docx(doc, "erd.png", "Figure 3.4: Entity-Relationship Diagram (ERD)", width_inches=6.0)
    
    doc.add_page_break()
    add_heading_2(doc, "3.6 Normalization (1NF to 3NF)")
    add_p(doc,
        "To avoid database design flaws (update, insertion, and deletion anomalies) and optimize storage, the MediFlow schema is normalized systematically from Unnormalized Form (UNF) through First (1NF), Second (2NF), and Third Normal Form (3NF)."
    )
    
    add_heading_3(doc, "3.6.1 First Normal Form (1NF)")
    add_p(doc,
        "A relation is in 1NF if all attributes contain only atomic (indivisible) values, and there are no repeating groups. In our design, all columns contain single values (e.g. no comma-separated text lists of multiple phone numbers or composite addresses in a single slot). Unique primary keys (patient_id, doctor_id, appointment_id, log_id, report_id) are defined to uniquely identify rows, thereby satisfying 1NF requirements."
    )
    
    add_heading_3(doc, "3.6.2 Second Normal Form (2NF)")
    add_p(doc,
        "A relation is in 2NF if it satisfies 1NF and contains no partial dependencies (i.e., no non-key attribute is dependent on only a subset of a composite primary key). Because all tables in MediFlow use single-attribute surrogate primary keys (e.g., patient_id, doctor_id, appointment_id) rather than composite keys, there is no possibility of partial dependencies. Every non-prime attribute (like name, email, specializing, date, status) depends fully on the whole single primary key. Thus, all tables satisfy 2NF."
    )
    
    add_heading_3(doc, "3.6.3 Third Normal Form (3NF)")
    add_p(doc,
        "A relation is in 3NF if it satisfies 2NF and contains no transitive dependencies (i.e. no non-key attribute is functionally dependent on another non-key attribute). Every non-key column in our tables depends directly and exclusively on the primary key, and nothing else (e.g., in patients table, email depends on patient_id, name depends on patient_id. In appointments table, status, reason, date, and time depend directly on appointment_id; the keys patient_id and doctor_id reference their respective primary keys without transitively linking names or specialties in the appointments table). Because no transitive dependencies exist, the schema satisfies 3NF."
    )

    # Chapter 4: Program Code
    doc.add_page_break()
    add_heading_1(doc, "Chapter 4: Program Code")
    
    add_heading_2(doc, "4.1 Code Details and Code Efficiency")
    add_p(doc,
        "The project code is optimized to enforce data processing logic directly at the database engine level wherever possible, yielding high code efficiency and minimal latency:"
    )
    add_bullet_point_docx(doc, "SQL View Efficiency", "By joining patients, doctors, and appointments inside the SQL View 'upcoming_appointments_view', the Express API server only performs a simple 'SELECT * FROM upcoming_appointments_view' query. This shifts join computation to MySQL's query optimizer and reduces Express server overhead.")
    add_bullet_point_docx(doc, "Trigger Autonomy", "The MySQL triggers run automatically on row insertions or updates. This ensures that even if developers execute manual SQL updates bypassing the Node.js Express API, the audit trail in appointment_logs remains intact and secure.")
    add_bullet_point_docx(doc, "Stored Procedure with Cursor", "The procedure 'generate_daily_report' uses a MySQL cursor to compile daily schedules. Running this cursor loop inside the database server eliminates the need to execute multiple SELECT roundtrips from Node.js, reducing network overhead.")
    
    # Load and write code contents
    schema_path = "database/schema.sql"
    server_path = "server.js"
    app_path = "public/app.js"
    
    try:
        with open(schema_path, "r", encoding="utf-8") as f:
            sql_code = f.read()
    except Exception as e:
        sql_code = f"Error reading {schema_path}: {str(e)}"
        
    try:
        with open(server_path, "r", encoding="utf-8") as f:
            server_code = f.read()
    except Exception as e:
        server_code = f"Error reading {server_path}: {str(e)}"
        
    try:
        with open(app_path, "r", encoding="utf-8") as f:
            app_lines = f.read().splitlines()
            app_code = "\n".join(app_lines[260:480])
    except Exception as e:
        app_code = f"Error reading {app_path}: {str(e)}"
        
    doc.add_page_break()
    add_heading_3(doc, "4.1.1 MySQL Database Schema Definition (schema.sql)")
    add_code_block_docx(doc, "database/schema.sql", sql_code)
    
    doc.add_page_break()
    add_heading_3(doc, "4.1.2 Node.js Express Backend API Router (server.js)")
    add_code_block_docx(doc, "server.js", server_code)
    
    doc.add_page_break()
    add_heading_3(doc, "4.1.3 Frontend Portal JavaScript Controllers (app.js - Excerpt)")
    add_code_block_docx(doc, "public/app.js (Selected API Integration Functions)", app_code)

    # Chapter 5: Result and Discussion
    doc.add_page_break()
    add_heading_1(doc, "Chapter 5: Result and Discussion")
    
    add_heading_2(doc, "5.1 User Documentation")
    add_p(doc,
        "The MediFlow application is designed with user ease-of-use in mind. To set up and run the system locally, follow these instructions:"
    )
    
    add_heading_3(doc, "5.1.1 Database Setup using XAMPP")
    add_bullet_point_docx(doc, "Step 1: Start MySQL in XAMPP", "Open XAMPP Control Panel and start MySQL on default port 3307 or 3306.")
    add_bullet_point_docx(doc, "Step 2: Install Node.js Dependencies", "Open a command prompt in the project root and run 'npm install' to resolve dependencies.")
    add_bullet_point_docx(doc, "Step 3: Auto-Setup Database", "Run 'node import-db.js'. This script connects to your running MySQL and imports database/schema.sql, inserting tables, trigger, procedures, and seed data automatically.")
    
    add_heading_3(doc, "5.1.2 Application Operations")
    add_bullet_point_docx(doc, "Start Server", "Run 'npm start' to start the Node.js server. The server runs at http://localhost:3000.")
    add_bullet_point_docx(doc, "Patient Portal", "Patients can access http://localhost:3000, register their profiles, select their profiles, view upcoming bookings, and book slots with specialist doctors.")
    add_bullet_point_docx(doc, "Doctor Portal", "Doctors select their name to load their queue of upcoming consultations, mark appointments 'Completed' or 'Cancelled', and run daily cursor stored procedure reports.")
    add_bullet_point_docx(doc, "Developer Lab Dashboard", "Developers can inspect raw database statistics (patient/doctor/appointment counts), watch database triggers log changes dynamically, and browse history logs.")

    # Chapter 6: Testing
    doc.add_page_break()
    add_heading_1(doc, "Chapter 6: Testing")
    
    add_heading_2(doc, "6.1 Unit Testing")
    add_p(doc,
        "Unit testing validates individual components and constraints to verify correct schema behavior:"
    )
    add_bullet_point_docx(doc, "Constraint Validation Test", "Assert that inserting a patient with a duplicate email fails with ER_DUP_ENTRY error code.")
    add_bullet_point_docx(doc, "Foreign Key Bounds Test", "Confirm that booking an appointment with a non-existent doctor_id (e.g. 99) fails with foreign key constraint errors.")
    add_bullet_point_docx(doc, "Enum Constraints Test", "Verify that appointments only accept 'Scheduled', 'Completed', and 'Cancelled' as statuses.")
    
    add_heading_2(doc, "6.2 Integration Testing")
    add_p(doc,
        "Integration testing validates functional pipelines, tracing flows between the client web page, backend server routes, and database objects:"
    )
    add_bullet_point_docx(doc, "Booking to View Integration Flow", "Register patient John -> Select John -> Submit Booking with Dr. Sarah -> Confirm that the client queries 'upcoming_appointments_view' and displays the new booking correctly.")
    add_bullet_point_docx(doc, "Trigger Integration Test", "Insert appointment -> Verify that the MySQL engine immediately executes trigger 'after_appointment_insert' and creates an audit row in 'appointment_logs'. Update status -> Confirm that 'after_appointment_update' records the state change details.")
    
    add_heading_2(doc, "6.3 System Testing")
    add_p(doc,
        "System testing evaluates the performance and concurrency limits under heavy load:"
    )
    add_bullet_point_docx(doc, "Connection Pool Verification", "Simulate multiple simultaneous requests using Apache Bench. Confirm that the Node.js mysql2 connection pool recycles database threads without timing out.")
    add_bullet_point_docx(doc, "Cascading Deletions Test", "Delete a patient profile -> Confirm that MySQL cascades the delete and automatically wipes all linked appointments for that patient, preventing orphan records.")
    
    add_heading_2(doc, "6.4 Acceptance Testing")
    add_p(doc,
        "Acceptance testing verifies that all project requirements specified by the user are satisfied:"
    )
    add_bullet_point_docx(doc, "User Acceptance Criteria", "Validate that the web interface is functional, responsive, and provides interactive tabs for patient profiles, doctor queues, database logs, and daily cursor reports.")
    add_bullet_point_docx(doc, "DBMS Features Checklist", "Verify that the three core elements (SQL View, SQL Triggers, and Cursor Stored Procedure) run successfully and integrate with the Node.js API endpoints.")

    # Chapter 7: Conclusion
    doc.add_page_break()
    add_heading_1(doc, "Chapter 7: Conclusion")
    
    add_heading_2(doc, "7.1 Conclusion & Future Enhancements")
    add_p(doc,
        "The development of the Patient Appointment Scheduling System ('MediFlow') demonstrates the effectiveness of relational databases in streamlining medical clinic workflows. By implementing constraints, views, triggers, and stored procedures directly at the database level, the system ensures data consistency, automates audit trails, and simplifies scheduling lookups. The 3-tier decoupling guarantees web responsiveness and clean server-client API integration."
    )
    add_p(doc,
        "Future enhancements could include:"
    )
    add_bullet_point_docx(doc, "Doctor Shift Rosters", "Develop a calendar module to restrict bookings to the doctor's specific working hours.")
    add_bullet_point_docx(doc, "Real-time Notifications", "Integrate SMS/Email API services (like Twilio or Nodemailer) to email patients their booking details.")
    add_bullet_point_docx(doc, "Patient Authentications", "Add JWT-based user login authentication to secure patient and doctor portals.")
    add_bullet_point_docx(doc, "Horizontal Clustering", "Set up database replication (Master-Slave) to allow high availability and backup support.")

    output_filename = "patient_appointment_system_report.docx"
    print(f"Saving compiled DOCX to {output_filename}...")
    try:
        doc.save(output_filename)
        print("Success: DOCX generation completed successfully!")
    except Exception as e:
        print(f"Error saving DOCX file: {str(e)}")

if __name__ == "__main__":
    build_docx_report()
